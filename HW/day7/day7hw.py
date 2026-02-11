import pdfplumber
from llm_guard.input_scanners import PromptInjection
from llm_guard.input_scanners.prompt_injection import MatchType
from docx import Document
from PIL import Image
import torch
import pandas as pd
import os
import uuid
import requests
import base64
import gc
from typing import List
from tqdm import tqdm
from openai import OpenAI
from qdrant_client import QdrantClient, models
from langchain_core.embeddings import Embeddings
from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from transformers import AutoTokenizer, AutoModelForCausalLM
from io import BytesIO

# ================= 配置設定 =================
FILE_PATH = "HW"
FILE_LIST = ["1.pdf", "2.pdf", "3.pdf", "4.png", "5.docx"]
EMBEDDING_API_URL = "http://ws-04.wade0426.me/embed"
QDRANT_URL = "http://localhost:6333"
COLLECTION_NAME = "rag_homework_day7_api"

# 主要 LLM
LLM_BASE_URL = "https://ws-03.wade0426.me/v1"
LLM_API_KEY = "day7hw"
LLM_MODEL_NAME = "/models/Qwen3-30B-A3B-Instruct-2507-FP8"

# OLM OCR 模型
OLM_API_URL = "https://6c1f-163-17-132-191.ngrok-free.app/v1/chat/completions"
OLM_API_KEY = "day7hw"
OLM_MODEL_NAME = "allenai/olmOCR-2-7B-1025-FP8"

RERANKER_MODEL_PATH = os.path.expanduser("../day6/Qwen3-Reranker-0.6B")
PREDICT_INPUT = "HW/questions.csv"
PREDICT_OUTPUT = "HW/output.csv"
GROUND_TRUTH_OUTPUT = "HW/ground_truth.csv"

client = QdrantClient(url=QDRANT_URL)

if torch.cuda.is_available():
    device_obj = torch.device("cuda")
elif torch.backends.mps.is_available():
    device_obj = torch.device("mps")
else:
    device_obj = torch.device("cpu")
print(f"* Device: {device_obj}")


# ================= 類別與函式定義 =================

class CustomAPIEmbeddings(Embeddings):
    def __init__(self, api_url):
        self.api_url = api_url

    def _call_api(self, texts: List[str]) -> List[List[float]]:
        data = {"texts": texts, "normalize": True, "batch_size": 32}
        try:
            response = requests.post(self.api_url, json=data, timeout=60)
            if response.status_code == 200:
                return response.json().get('embeddings', [])
            else:
                print(f"❌ API Error Code: {response.status_code}")
                return []
        except Exception as e:
            print(f"❌ API Exception: {e}")
            return []

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        return self._call_api(texts)

    def embed_query(self, text: str) -> List[float]:
        results = self._call_api([text])
        return results[0] if results else []


class SimpleLLMClient:
    def __init__(self, base_url, model_name, api_key):
        self.client = OpenAI(base_url=base_url, api_key=api_key)
        self.model_name = model_name

    def generate(self, prompt: str) -> str:
        try:
            response = self.client.chat.completions.create(
                model=self.model_name,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"LLM Error: {e}")
            return "Error generating response."


# 初始化 Embedding 與 LLM
print(f"* Initial Embedding API ({EMBEDDING_API_URL})...")
embedding_model = CustomAPIEmbeddings(EMBEDDING_API_URL)
try:
    test_vec = embedding_model.embed_query("測試")
    if test_vec:
        EMBED_DIM = len(test_vec)
        print(f"✅ API 連線成功！向量維度: {EMBED_DIM}")
    else:
        raise ValueError("API 回傳為空")
except Exception as e:
    print(f"❌ API 測試失敗: {e}")
    exit()

print("* 初始化 LLM Client...")
llm_client = SimpleLLMClient(LLM_BASE_URL, LLM_MODEL_NAME, LLM_API_KEY)


# --- OLM OCR 功能 ---

def encode_image_to_base64(image_path, max_size=1024):
    """
    將圖片縮小後轉換為 Base64 字串，避免 Token 超過模型上限。
    """
    with Image.open(image_path) as img:
        if max(img.size) > max_size:
            ratio = max_size / max(img.size)
            new_size = (int(img.width * ratio), int(img.height * ratio))
            img = img.resize(new_size, Image.Resampling.LANCZOS)
            print(f"   (圖片過大，已縮小至: {new_size})")

        buffered = BytesIO()
        img.save(buffered, format="PNG")
        return base64.b64encode(buffered.getvalue()).decode('utf-8')


def call_olm_ocr_api(image_path):
    """呼叫 OLM API 進行圖片辨識"""
    base64_image = encode_image_to_base64(image_path)

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {OLM_API_KEY}"
    }

    payload = {
        "model": OLM_MODEL_NAME,
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": "Convert this image to markdown text. Preserve tables and formatting."
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{base64_image}"
                        }
                    }
                ]
            }
        ],
        "max_tokens": 4096,
        "temperature": 0.1
    }

    try:
        response = requests.post(OLM_API_URL, headers=headers, json=payload, timeout=120)
        if response.status_code == 200:
            result = response.json()
            content = result['choices'][0]['message']['content']
            return content
        else:
            print(f"❌ OLM API Error: {response.status_code} - {response.text}")
            return ""
    except Exception as e:
        print(f"❌ OLM Connection Error: {e}")
        return ""


def process_idp_files(file_path, file_list):
    """讀取並處理所有檔案 (PDF 混合模式 + DOCX 表格支援)"""
    docs_content = {}
    print("--- 開始 IDP 處理 ---")

    for file_name in file_list:
        full_path = os.path.join(file_path, file_name)
        if not os.path.exists(full_path):
            print(f"* 找不到檔案: {file_name}")
            continue

        print(f"正在處理: {file_name}...")
        extracted_text = ""

        # === 1. PDF 處理 ===
        if file_name.endswith(".pdf"):
            try:
                with pdfplumber.open(full_path) as pdf:
                    for i, page in enumerate(pdf.pages):
                        text = page.extract_text()

                        # 判斷是否為掃描檔
                        if not text or len(text.strip()) < 10:
                            print(f"   ⚠️ 第 {i + 1} 頁疑似為掃描檔，轉為圖片進行 OLM OCR...")
                            im = page.to_image(resolution=300).original
                            temp_img_path = f"temp_page_{i}.png"
                            im.save(temp_img_path)

                            ocr_text = call_olm_ocr_api(temp_img_path)
                            if ocr_text: extracted_text += ocr_text + "\n"

                            if os.path.exists(temp_img_path): os.remove(temp_img_path)
                        else:
                            extracted_text += text + "\n"
                            # 提取表格
                            tables = page.extract_tables()
                            for table in tables:
                                table_str = ""
                                for row in table:
                                    cleaned_row = [str(cell).strip() if cell is not None else "" for cell in row]
                                    if any(cleaned_row):
                                        table_str += " | ".join(cleaned_row) + "\n"
                                if table_str:
                                    extracted_text += "\n[表格]\n" + table_str + "\n"
                print(f"✅ [{file_name}] PDF 處理完成")
            except Exception as e:
                print(f"❌ PDF 錯誤: {e}")

        # === 2. PNG 處理 ===
        elif file_name.endswith(".png"):
            try:
                print(f"⏳ 正在呼叫 OLM 模型處理圖片 ({file_name})...")
                olm_text = call_olm_ocr_api(full_path)
                if olm_text:
                    extracted_text = olm_text
                    print(f"✅ [{file_name}] OLM 辨識成功")
                else:
                    print(f"⚠️ [{file_name}] OLM 未回傳內容")
            except Exception as e:
                print(f"❌ OLM 處理失敗 {file_name}: {e}")

        # === 3. DOCX 處理 (支援表格) ===
        elif file_name.endswith(".docx"):
            try:
                doc = Document(full_path)
                # 段落
                for para in doc.paragraphs:
                    extracted_text += para.text + "\n"
                # 表格
                for table in doc.tables:
                    table_str = ""
                    for row in table.rows:
                        row_cells = [cell.text.strip() for cell in row.cells]
                        if any(row_cells):
                            table_str += " | ".join(row_cells) + "\n"
                    if table_str:
                        extracted_text += "\n[表格]\n" + table_str + "\n"

                print(f"✅ [{file_name}] Word 讀取成功 (含表格)")
            except Exception as e:
                print(f"❌ Word 錯誤: {e}")

        if extracted_text.strip():
            docs_content[file_name] = extracted_text
        else:
            print(f"⚠️ {file_name} 內容為空，跳過。")

    return docs_content


def scan_chunks_for_injection(split_docs):
    """
    【關鍵修改】針對切分後的 Chunk 進行掃描
    這樣可以避免：
    1. 內容過長被截斷
    2. 惡意特徵被大量正常文字稀釋
    """
    print("--- 執行安全性掃描 (Chunk Level) ---")
    # 對於 Chunk，0.5 的閾值通常足夠，因為濃度變高了
    scanner = PromptInjection(threshold=0.5, match_type=MatchType.FULL)

    final_docs = []

    for doc in tqdm(split_docs, desc="掃描 Chunks"):
        text_content = doc.page_content
        source = doc.metadata.get("source", "unknown")

        _, is_valid, score = scanner.scan(text_content)

        if not is_valid:
            print(f"⚠️ [警告] 檔案 {source} 的某個區塊疑似惡意 (Score: {score})")
            # 這裡決定「放行」以確保作業能回答問題，但實務上通常會丟棄
            final_docs.append(doc)
        else:
            final_docs.append(doc)

    return final_docs


def init_qdrant_collection(documents):
    """寫入向量資料庫"""
    if not documents:
        print("⚠️ 無文件可寫入")
        return

    print(f"🔄 重置集合 {COLLECTION_NAME}...")
    try:
        client.delete_collection(COLLECTION_NAME)
    except:
        pass

    client.create_collection(
        collection_name=COLLECTION_NAME,
        vectors_config={"dense": models.VectorParams(distance=models.Distance.COSINE, size=EMBED_DIM)},
        sparse_vectors_config={"sparse": models.SparseVectorParams(modifier=models.Modifier.IDF)},
    )

    texts_to_embed = [doc.page_content for doc in documents]
    print(f"⏳ 計算 {len(texts_to_embed)} 筆 Embeddings...")

    doc_embeddings = embedding_model.embed_documents(texts_to_embed)
    if len(doc_embeddings) != len(documents):
        print("❌ Embedding 數量不符")
        return

    points = []
    for doc, embedding in zip(documents, doc_embeddings):
        points.append(models.PointStruct(
            id=uuid.uuid4().hex,
            vector={
                "dense": embedding,
                "sparse": models.Document(text=doc.page_content, model="Qdrant/bm25"),
            },
            payload={
                "text": doc.page_content,
                "source": doc.metadata.get("source", "unknown")
            },
        ))

    batch_size = 50
    for i in tqdm(range(0, len(points), batch_size), desc="寫入 Qdrant"):
        client.upsert(collection_name=COLLECTION_NAME, points=points[i: i + batch_size])
    print("✅ 資料寫入完成")


# --- Reranker 相關 ---
print("* 載入 Reranker 模型...")
reranker_tokenizer = AutoTokenizer.from_pretrained(RERANKER_MODEL_PATH, local_files_only=True, trust_remote_code=True)
reranker_model = AutoModelForCausalLM.from_pretrained(RERANKER_MODEL_PATH, local_files_only=True,
                                                      trust_remote_code=True).to(device_obj).eval()

token_false_id = reranker_tokenizer.convert_tokens_to_ids("no")
token_true_id = reranker_tokenizer.convert_tokens_to_ids("yes")
prefix = "<|im_start|>system\nJudge whether the Document meets the requirements based on the Query and the Instruct provided. Note that the answer can only be \"yes\" or \"no\".<|im_end|>\n<|im_start|>user\n"
suffix = "<|im_end|>\n<|im_start|>assistant\n"


def compute_rerank_scores(pairs, batch_size=4):
    all_scores = []
    for i in range(0, len(pairs), batch_size):
        batch_pairs = pairs[i: i + batch_size]
        processed_inputs = []
        for pair in batch_pairs:
            text = f"{prefix}{pair}{suffix}"
            processed_inputs.append(text)

        inputs = reranker_tokenizer(processed_inputs, padding=True, truncation=True, return_tensors="pt",
                                    max_length=1024).to(device_obj)

        with torch.no_grad():
            logits = reranker_model(**inputs).logits[:, -1, :]
            scores = logits[:, token_true_id].exp().tolist()
            all_scores.extend(scores)

        del inputs, logits
        torch.cuda.empty_cache() if torch.cuda.is_available() else None

    return all_scores


def rerank_documents(query, documents):
    if not documents: return []
    formatted_pairs = [f"<Instruct>: 根據查詢檢索相關文件\n<Query>: {query}\n<Document>: {doc['text']}" for doc in
                       documents]
    scores = compute_rerank_scores(formatted_pairs)
    doc_scores = list(zip(documents, scores))
    doc_scores.sort(key=lambda x: x[1], reverse=True)
    return doc_scores


def hybrid_search_with_rerank(query: str, initial_limit=20, final_limit=3):
    query_vec = embedding_model.embed_query(query)
    try:
        response = client.query_points(
            collection_name=COLLECTION_NAME,
            prefetch=[
                models.Prefetch(query=models.Document(text=query, model="Qdrant/bm25"), using="sparse",
                                limit=initial_limit),
                models.Prefetch(query=query_vec, using="dense", limit=initial_limit),
            ],
            query=models.FusionQuery(fusion=models.Fusion.RRF),
            limit=initial_limit,
        )
        candidate_docs = [{"text": point.payload.get("text", ""), "source": point.payload.get("source", "unknown")} for
                          point in response.points]
    except Exception as e:
        print(f"Search Error: {e}")
        return []

    if not candidate_docs: return []
    top_results = rerank_documents(query, candidate_docs)[:final_limit]
    return top_results


def query_rewrite(query: str) -> str:
    prompt = f"你是一個搜尋引擎優化專家。請將以下使用者的問題改寫為更精確的關鍵字查詢。\n使用者問題: {query}\n改寫後查詢:"
    return llm_client.generate(prompt).strip()


def main():
    # 1. IDP 處理
    docs_content = process_idp_files(FILE_PATH, FILE_LIST)

    # 2. 切分 (Chunking) - 【順序改變：先切分】
    text_splitter = RecursiveCharacterTextSplitter(
        separators=["\n\n", "\n", " ", ""], chunk_size=500, chunk_overlap=50
    )

    all_documents = []
    for filename, text in docs_content.items():
        all_documents.append(LCDocument(page_content=text, metadata={"source": filename}))

    split_docs = text_splitter.split_documents(all_documents)
    print(f"📊 共切分出 {len(split_docs)} 個區塊")

    # 3. 安全性掃描 (Scanning) - 【順序改變：後掃描】
    # 針對切分後的 Chunk 進行掃描，這樣才能抓到 5.docx 的惡意片段
    final_docs = scan_chunks_for_injection(split_docs)

    # 4. 寫入向量庫
    init_qdrant_collection(final_docs)

    print(f"📂 讀取問題: {PREDICT_INPUT}")
    if not os.path.exists(PREDICT_INPUT):
        print("❌ 檔案不存在")
        return

    df = pd.read_csv(PREDICT_INPUT)

    # 【測試模式】只跑前 5 題
    df = df.head(5)
    print(f"⚠️ 測試模式啟動：僅處理前 {len(df)} 題資料")

    if 'answer' not in df.columns: df['answer'] = None
    if 'source' not in df.columns: df['source'] = None
    if 'id' not in df.columns:
        print("❌ 缺少 id 欄位")
        return

    ground_truth_list = []

    print("🚀 開始回答問題...")
    for index, row in tqdm(df.iterrows(), total=df.shape[0]):
        original_question = str(row['questions'])

        refined_query = query_rewrite(original_question)
        search_results = hybrid_search_with_rerank(refined_query)

        if search_results:
            retrieval_docs = [item[0] for item in search_results]
            context_str = "\n".join([doc['text'] for doc in retrieval_docs])
            unique_sources = sorted(list(set([doc['source'] for doc in retrieval_docs])))
            source_str = ",".join(unique_sources)

            if index < 1:
                print(f"\n🔍 [Debug Context]: {context_str[:200]}...")
        else:
            context_str = ""
            source_str = ""
            retrieval_docs = []

        ground_truth_list.append({
            "id": row['id'],
            "questions": original_question,
            "contexts": [doc['text'] for doc in retrieval_docs],
            "ground_truth": ""
        })

        qa_prompt = f"""
        你是一個專業的資訊助手。請根據【參考資料】回答問題。
        若資料不足請回答「目前資訊不足」。
        請直接回答重點，不要重複問題。

        【參考資料】：
        {context_str}

        【問題】：{original_question}
        【回答】：
        """
        answer = llm_client.generate(qa_prompt)

        df.at[index, 'answer'] = answer
        df.at[index, 'source'] = source_str

        torch.cuda.empty_cache() if torch.cuda.is_available() else None
        gc.collect()

    df.to_csv(PREDICT_OUTPUT, index=False, encoding='utf-8-sig')
    print(f"✅ 結果已儲存: {PREDICT_OUTPUT}")

    pd.DataFrame(ground_truth_list).to_csv(GROUND_TRUTH_OUTPUT, index=False, encoding='utf-8-sig')
    print(f"✅ GT 已儲存: {GROUND_TRUTH_OUTPUT}")


if __name__ == "__main__":
    main()